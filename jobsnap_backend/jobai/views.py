from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
import json
from docx import Document
from io import BytesIO

@csrf_exempt
@require_POST
def generate_reply(request):
    try:
        data = json.loads(request.body)
        prompt = data.get("prompt", "")
        tone = data.get("tone", "")
        reply = f"Replying in {tone} tone to: {prompt}"

        return JsonResponse({
            "status": "success",
            "reply": reply
        })

    except Exception as e:
        return JsonResponse({
            "status": "error",
            "message": str(e)
        }, status=400)

@csrf_exempt
@require_POST
def generate_resume(request):
    try:
        data = json.loads(request.body)
        name = data.get("name", "Unnamed")
        skills = data.get("skills", [])
        experience = data.get("experience", "")

        # Create DOCX document
        doc = Document()
        doc.add_heading(f"Resume - {name}", 0)

        doc.add_heading("Name", level=1)
        doc.add_paragraph(name)

        doc.add_heading("Skills", level=1)
        for skill in skills:
            doc.add_paragraph(f"• {skill}", style='List Bullet')

        doc.add_heading("Experience", level=1)
        doc.add_paragraph(experience)

        # Save to memory
        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)

        response = HttpResponse(
            file_stream.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename={name}_resume.docx'
        return response

    except Exception as e:
        return JsonResponse({
            "status": "error",
            "message": str(e)
        }, status=400
)